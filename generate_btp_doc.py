from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

SAP_BLUE   = RGBColor(0x00, 0x33, 0x66)
MID_BLUE   = RGBColor(0x00, 0x5A, 0x9C)
DARK_GREY  = RGBColor(0x33, 0x33, 0x33)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
GREEN      = RGBColor(0x16, 0xA3, 0x4A)
AMBER      = RGBColor(0xD9, 0x77, 0x06)
RED        = RGBColor(0xDC, 0x26, 0x26)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="DDDDDD"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(16); run.font.bold = True; run.font.color.rgb = SAP_BLUE
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1"); bot.set(qn("w:color"), "003366")
    pBdr.append(bot); pPr.append(pBdr)

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(13); run.font.bold = True; run.font.color.rgb = MID_BLUE

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11); run.font.bold = True; run.font.color.rgb = DARK_GREY

def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(10.5); r.font.color.rgb = DARK_GREY

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.8 + level * 0.6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(10.5); r.font.color.rgb = DARK_GREY

def code_block(lines):
    p = doc.add_paragraph()
    r = p.add_run("\n".join(lines))
    r.font.name = "Courier New"; r.font.size = Pt(8.5); r.font.color.rgb = DARK_GREY
    doc.add_paragraph()

def info_box(title, lines, color_hex="E8F0FE", border_hex="003366"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, color_hex)
    set_cell_borders(cell, border_hex)
    cell.paragraphs[0].clear()
    hr = cell.paragraphs[0].add_run(title)
    hr.font.bold = True; hr.font.size = Pt(10.5); hr.font.color.rgb = SAP_BLUE
    for line in lines:
        p = cell.add_paragraph()
        r = p.add_run(line)
        r.font.size = Pt(10); r.font.color.rgb = DARK_GREY
    doc.add_paragraph()

def data_table(headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, "003366")
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(h)
        r.font.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row in enumerate(rows):
        tr = tbl.add_row()
        bg = "F5F5F5" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            set_cell_borders(cell, "DDDDDD")
            cell.paragraphs[0].clear()
            r = cell.paragraphs[0].add_run(val)
            r.font.size = Pt(10); r.font.color.rgb = DARK_GREY
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()

# ── COVER PAGE ────────────────────────────────────────────────────────────────
banner = doc.add_table(rows=1, cols=1)
banner.style = "Table Grid"
bc = banner.cell(0, 0)
set_cell_bg(bc, "003366")
bc.paragraphs[0].clear()
bc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
br1 = bc.paragraphs[0].add_run("\n\nLTVF Cloud Dashboard\n")
br1.font.size = Pt(28); br1.font.bold = True; br1.font.color.rgb = WHITE
br2 = bc.add_paragraph()
br2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = br2.add_run("SAP BTP Integration — Technical Design Document\n\n")
r2.font.size = Pt(14); r2.font.color.rgb = RGBColor(0xB3, 0xD1, 0xFF)
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in [
    "Integration Approach: SAP BTP Destination + Cloud Connector (On-Premise)",
    f"Date: {datetime.date.today().strftime('%d %B %Y')}",
    "Version: 1.0",
    "BTP Global Account: DMLT BTP Global Account (Canary)",
    "BTP Subaccount: BDV space  |  85de1f19-b01f-4eb6-b07b-8037a9d9431a",
]:
    r = meta.add_run(line + "\n")
    r.font.size = Pt(10.5); r.font.color.rgb = DARK_GREY

doc.add_page_break()

# ── 1. OVERVIEW ───────────────────────────────────────────────────────────────
heading1("1.  Overview")
body(
    "This document describes the SAP BTP integration added to the LTVF Cloud Dashboard. "
    "The integration enables users to fetch live CNVLTVF3 test results directly from the "
    "on-premise SAP system via a 'Fetch from SAP BTP' button on the upload screen — "
    "eliminating the need to manually export an Excel file."
)
body(
    "The existing Excel upload flow is fully preserved. The SAP BTP fetch is an additional "
    "option that appears automatically on the upload screen when the backend is configured."
)

info_box("BTP Account Details", [
    "  Global Account:   DMLT BTP Global Account (Canary)",
    "  Global Account ID: 79b39184-de86-4342-8a71-7d4b98c7712f",
    "  Subaccount:       BDV space",
    "  Subaccount ID:    85de1f19-b01f-4eb6-b07b-8037a9d9431a",
    "  Environment:      Canary  (canary.cockpit.btp.int.sap)",
])

doc.add_page_break()

# ── 2. ARCHITECTURE ───────────────────────────────────────────────────────────
heading1("2.  Architecture")

heading2("2.1  Integration Flow")
code_block([
    "┌─────────────────────────────────────────────────────────────────┐",
    "│  USER BROWSER                                                   │",
    "│  React SPA — clicks 'Fetch from SAP BTP' button                │",
    "└──────────────────────────┬──────────────────────────────────────┘",
    "                           │  GET /api/sap/fetch  (HTTPS)",
    "┌──────────────────────────▼──────────────────────────────────────┐",
    "│  FASTAPI BACKEND  (Render.com)                                  │",
    "│                                                                 │",
    "│  Step 1: POST BTP_TOKEN_URL → OAuth2 access_token (XSUAA)      │",
    "│  Step 2: GET  BTP_DEST_SVC_URL/.../{LTVF_ONPREMISE}            │",
    "│          → resolves virtual host, auth headers, proxy info      │",
    "│  Step 3: GET  OData via BTP Connectivity proxy                  │",
    "│          → /sap/opu/odata/sap/CNVLTVF3_SRV/LTVFResultSet       │",
    "│  Step 4: Map OData JSON → LTVFParseResult → return to browser   │",
    "└──────────────────────────┬──────────────────────────────────────┘",
    "            OAuth2 + proxy │  HTTPS through BTP Connectivity Service",
    "┌──────────────────────────▼──────────────────────────────────────┐",
    "│  SAP BTP — BDV space subaccount                                 │",
    "│  ┌─────────────────────┐   ┌──────────────────────────────┐    │",
    "│  │  XSUAA (OAuth2)     │   │  Destination Service         │    │",
    "│  │  issues tokens      │   │  resolves LTVF_ONPREMISE     │    │",
    "│  └─────────────────────┘   └──────────────────────────────┘    │",
    "│  ┌──────────────────────────────────────────────────────────┐  │",
    "│  │  Connectivity Service — proxies to Cloud Connector tunnel │  │",
    "│  └──────────────────────────────┬───────────────────────────┘  │",
    "└─────────────────────────────────┼───────────────────────────────┘",
    "                      SCC tunnel  │  (outbound only, no inbound FW)",
    "┌─────────────────────────────────▼───────────────────────────────┐",
    "│  CORPORATE NETWORK                                              │",
    "│  SAP Cloud Connector  →  On-Premise SAP System                  │",
    "│  Virtual: sap-ltvf:8000  →  Real: <sap-internal-host>:<port>   │",
    "└─────────────────────────────────────────────────────────────────┘",
])

heading2("2.2  Three BTP Services Used")
data_table(
    ["BTP Service", "Purpose", "Credentials needed"],
    [
        ["XSUAA (OAuth2)",        "Issues access tokens for authenticating API calls to BTP services",           "BTP_TOKEN_URL (from Destination Service key → url + /oauth/token)"],
        ["Destination Service",   "Resolves named destination LTVF_ONPREMISE to URL, auth headers, proxy info", "BTP_DEST_CLIENT_ID, BTP_DEST_CLIENT_SECRET, BTP_DEST_SVC_URL"],
        ["Connectivity Service",  "HTTP proxy that routes requests through the Cloud Connector tunnel",          "BTP_CONN_PROXY_HOST, BTP_CONN_PROXY_PORT (from Connectivity service key)"],
    ],
    col_widths=[1.6, 2.8, 2.8]
)

heading2("2.3  Files Added / Modified")
data_table(
    ["File", "Action", "Description"],
    [
        ["backend/btp_client.py",          "New",  "BTP OAuth2 token fetch, Destination Service lookup, OData call via proxy, OData→LTVFParseResult mapper"],
        ["backend/main.py",                "Edit", "Added GET /api/sap/status and GET /api/sap/fetch routes importing from btp_client.py"],
        ["backend/.env.example",           "Edit", "Added all BTP_* env var entries with step-by-step setup instructions"],
        ["backend/.env",                   "New",  "Local env file (git-ignored) — BTP_* values to be filled in once service instances are created"],
        ["frontend/src/api/btpApi.ts",     "New",  "checkSAPStatus() and fetchFromSAP() API functions"],
        ["frontend/src/App.tsx",           "Edit", "sapAvailable state, useEffect status check on mount, handleFetchFromSAP handler, props passed to UploadZone"],
        ["frontend/src/components/UploadZone.tsx", "Edit", "Fetch from SAP BTP button with Live badge — shown only when sapAvailable is true"],
    ],
    col_widths=[2.5, 0.7, 3.5]
)

doc.add_page_break()

# ── 3. BACKEND ────────────────────────────────────────────────────────────────
heading1("3.  Backend Implementation")

heading2("3.1  New API Routes (main.py)")
data_table(
    ["Method", "Route", "Behaviour"],
    [
        ["GET", "/api/sap/status", 'Returns {"available": true/false, "mode": "btp"}. Frontend calls this on mount to decide whether to show the Fetch button.'],
        ["GET", "/api/sap/fetch",  "Calls fetch_ltvf_via_btp(). Returns LTVFParseResult JSON on success. HTTP 503 if BTP not configured; HTTP 502 on SAP call failure with traceback."],
    ],
    col_widths=[0.7, 1.8, 4.2]
)

heading2("3.2  btp_client.py — Function Reference")
data_table(
    ["Function", "Description"],
    [
        ["is_btp_configured()",          "Checks that BTP_TOKEN_URL, BTP_DEST_CLIENT_ID, BTP_DEST_CLIENT_SECRET, BTP_DEST_SVC_URL are all non-empty. Used by status route."],
        ["_get_token(url, id, secret)",  "POSTs to XSUAA with grant_type=client_credentials. Returns bearer token string. Timeout: 15s."],
        ["_get_destination(dest_name)",  "Authenticates to Destination Service, calls /destination-configuration/v1/destinations/{name}. Returns full destination config dict."],
        ["_map_odata_to_result(json)",   "Maps OData d.results[] to LTVFParseResult. Field names controlled by FIELD_* constants. Root node (level 0, no parent) used for grand totals."],
        ["fetch_ltvf_via_btp()",         "Orchestrates: get destination → extract URL/proxy/auth → call OData → map result. This is the only public function called by main.py."],
    ],
    col_widths=[2.2, 4.5]
)

heading2("3.3  OData Field Mapping")
body(
    "The FIELD_* constants at the top of btp_client.py control which OData field names are "
    "used. Adjust these if BASIS exposes different field names in CNVLTVF3_SRV — no logic "
    "changes needed."
)
data_table(
    ["Constant", "Default OData field", "Maps to LTVFRow field"],
    [
        ["FIELD_NODE_ID",    "NodeId",         "id"],
        ["FIELD_PARENT_ID",  "ParentNodeId",   "parent_id"],
        ["FIELD_LEVEL",      "HierarchyLevel", "level"],
        ["FIELD_TEST_NAME",  "Description",    "test_name"],
        ["FIELD_IS_GROUP",   "IsGroup",        "is_group"],
        ["FIELD_RATE_PCT",   "MatchRate",      "rate_pct"],
        ["FIELD_DIFF",       "DiffCount",      "diff"],
        ["FIELD_MISSING",    "MissingCount",   "missing"],
        ["FIELD_UNEXPECTED", "UnexpectedCount","unexpected"],
        ["FIELD_EQUAL",      "EqualCount",     "equal"],
        ["FIELD_SOURCE",     "SourceVolume",   "source"],
        ["FIELD_TARGET",     "TargetVolume",   "target"],
        ["ODATA_ENTITY_SET", "LTVFResultSet",  "(entity set name in URL)"],
    ],
    col_widths=[2.0, 2.0, 2.0]
)

heading2("3.4  Environment Variables (backend/.env)")
data_table(
    ["Variable", "Source", "Description"],
    [
        ["BTP_TOKEN_URL",          "Destination Service key → url + /oauth/token",        "XSUAA token endpoint"],
        ["BTP_DEST_CLIENT_ID",     "Destination Service key → clientid",                  "OAuth2 client ID for Destination Service"],
        ["BTP_DEST_CLIENT_SECRET", "Destination Service key → clientsecret",              "OAuth2 client secret for Destination Service"],
        ["BTP_DEST_SVC_URL",       "Destination Service key → uri",                       "Destination Service REST API base URL"],
        ["BTP_CONN_PROXY_HOST",    "Connectivity Service key → onpremise_proxy_host",     "BTP Connectivity proxy hostname"],
        ["BTP_CONN_PROXY_PORT",    "Connectivity Service key → onpremise_proxy_port",     "BTP Connectivity proxy port (default 20003)"],
        ["SAP_DESTINATION_NAME",   "Name you choose when creating the BTP Destination",   "Name of destination pointing to on-premise SAP (default: LTVF_ONPREMISE)"],
    ],
    col_widths=[2.2, 2.5, 2.0]
)

doc.add_page_break()

# ── 4. FRONTEND ───────────────────────────────────────────────────────────────
heading1("4.  Frontend Implementation")

heading2("4.1  btpApi.ts")
body("Two functions mirror the two new backend routes:")
code_block([
    "// frontend/src/api/btpApi.ts",
    "",
    "checkSAPStatus(): Promise<{ available: boolean }>",
    "  → GET /api/sap/status",
    "  → called on App mount via useEffect",
    "  → result stored in sapAvailable state",
    "",
    "fetchFromSAP(): Promise<LTVFParseResult>",
    "  → GET /api/sap/fetch",
    "  → called when user clicks 'Fetch from SAP BTP'",
    "  → returns same LTVFParseResult as Excel upload",
])

heading2("4.2  App.tsx Changes")
data_table(
    ["Change", "Detail"],
    [
        ["sapAvailable state",     "boolean — stores result of checkSAPStatus(). Initialised false. Passed as prop to UploadZone."],
        ["useEffect on mount",     "Calls checkSAPStatus() once on load. Sets sapAvailable=true if backend has BTP configured. Silently catches errors (button stays hidden)."],
        ["handleFetchFromSAP()",   "Identical flow to handleFile(): sets loading=true → calls fetchFromSAP() → setData() → saveToHistory() → setTab('overview'). Error written to error state."],
        ["UploadZone props",       "sapAvailable and onFetchFromSAP passed to UploadZone. No other components changed."],
    ],
    col_widths=[2.0, 4.7]
)

heading2("4.3  UploadZone.tsx Changes")
body(
    "Two new optional props added: sapAvailable (boolean, default false) and "
    "onFetchFromSAP (callback). When sapAvailable is true, the component renders "
    "an 'or' divider and a 'Fetch from SAP BTP' button below the drag-and-drop zone."
)
body("The button is:")
for item in [
    "Full width (480px) matching the upload zone",
    "SAP blue background (#003366) with hover darkening",
    "CloudDownload icon (Lucide React — already in dependencies)",
    "Live badge (small pill, right-aligned)",
    "Disabled with reduced opacity when loading is true",
    "Hidden entirely when sapAvailable is false — zero visual change for unconfigured deployments",
]:
    bullet(item)

doc.add_page_break()

# ── 5. PENDING SETUP STEPS ────────────────────────────────────────────────────
heading1("5.  Pending Setup Steps")

body(
    "The code is complete. The following external steps are required before the "
    "'Fetch from SAP BTP' button becomes active in the deployed dashboard."
)

heading2("5.1  Step-by-Step Setup Checklist")
data_table(
    ["#", "Action", "Who", "Status"],
    [
        ["1", "Assign Subaccount Administrator + Connectivity and Destination Administrator role collections to user HV in BDV space subaccount", "BTP Global Account Admin", "Pending"],
        ["2", "Install SAP Cloud Connector on a server inside the corporate network", "BASIS / Infrastructure", "Pending"],
        ["3", "Connect Cloud Connector to BDV space subaccount (using Download Authentication Data from BTP Cockpit)", "BASIS / Infrastructure", "Pending"],
        ["4", "Expose on-premise SAP as virtual host sap-ltvf:8000 in Cloud Connector admin UI + add resource /sap/opu/odata/sap/CNVLTVF3_SRV/", "BASIS / Infrastructure", "Pending"],
        ["5", "Create read-only SAP service account with CNVLTVF3 + CNVLTVF3_SRV authorisation", "BASIS team", "Pending"],
        ["6", "Create Destination Service instance (plan: lite) in BDV space → download service key", "Hariprasad (after step 1)", "Pending"],
        ["7", "Create Connectivity Service instance (plan: lite) in BDV space → download service key", "Hariprasad (after step 1)", "Pending"],
        ["8", "Create named Destination LTVF_ONPREMISE in BTP Cockpit pointing to virtual host sap-ltvf:8000 with SAP service account credentials", "Hariprasad (after steps 1,4,5)", "Pending"],
        ["9", "Fill BTP_* values in backend/.env using service key credentials", "Hariprasad (after steps 6,7)", "Pending"],
        ["10","Set BTP_* env vars in Render.com backend service → redeploy", "Hariprasad (after step 9)", "Pending"],
    ],
    col_widths=[0.3, 3.2, 1.8, 1.0]
)

heading2("5.2  BTP Destination Configuration (step 8)")
body("Create a destination in BTP Cockpit → Connectivity → Destinations → New Destination:")
data_table(
    ["Field", "Value"],
    [
        ["Name",            "LTVF_ONPREMISE"],
        ["Type",            "HTTP"],
        ["URL",             "http://sap-ltvf:8000  (virtual host set in Cloud Connector)"],
        ["Proxy Type",      "OnPremise"],
        ["Authentication",  "BasicAuthentication"],
        ["User",            "<SAP service account username from BASIS>"],
        ["Password",        "<SAP service account password from BASIS>"],
        ["sap-client",      "100  (Additional Property)"],
        ["HTML5.DynamicDestination", "true  (Additional Property)"],
    ],
    col_widths=[2.2, 4.5]
)

heading2("5.3  .env Values to Fill In (step 9)")
code_block([
    "# backend/.env",
    "",
    "# From Destination Service instance service key:",
    "BTP_TOKEN_URL=https://<subdomain>.authentication.sap.hana.ondemand.com/oauth/token",
    "BTP_DEST_CLIENT_ID=sb-xsappname...",
    "BTP_DEST_CLIENT_SECRET=<clientsecret>",
    "BTP_DEST_SVC_URL=https://destination-configuration.cfapps.sap.hana.ondemand.com",
    "",
    "# From Connectivity Service instance service key:",
    "BTP_CONN_PROXY_HOST=connectivity.cf.sap.hana.ondemand.com",
    "BTP_CONN_PROXY_PORT=20003",
    "",
    "# Destination name (matches what you created in BTP Cockpit):",
    "SAP_DESTINATION_NAME=LTVF_ONPREMISE",
])

doc.add_page_break()

# ── 6. SECURITY ───────────────────────────────────────────────────────────────
heading1("6.  Security Design")

data_table(
    ["Control", "Implementation"],
    [
        ["OAuth2 token-based auth",       "Backend authenticates to BTP services using client_credentials flow. Tokens are short-lived and fetched per request."],
        ["No credentials in browser",     "All BTP credentials (client_id, client_secret) are backend-only env vars. The React frontend never sees them."],
        ["Cloud Connector tunnel",        "SAP is never directly exposed to the internet. All traffic flows through the encrypted BTP Connectivity tunnel via the Cloud Connector."],
        ["Read-only SAP service account", "The SAP account used in the LTVF_ONPREMISE destination has read-only CNVLTVF3 access only — no write risk."],
        ["CORS policy",                   "Backend ALLOWED_ORIGINS env var restricts /api/sap/* to the known frontend origin only."],
        [".env git-ignored",              ".gitignore already excludes .env — BTP credentials are never committed to the repository."],
        ["HTTPS everywhere",              "BTP token URL, Destination Service URL, and Render.com all enforce HTTPS. No credentials travel in plaintext."],
    ],
    col_widths=[2.2, 4.5]
)

doc.add_page_break()

# ── 7. TESTING & VERIFICATION ─────────────────────────────────────────────────
heading1("7.  Testing & Verification")

heading2("7.1  Backend — Before BTP is configured")
for step in [
    "Start backend: uvicorn main:app --reload",
    "GET http://localhost:8000/api/sap/status  →  {\"available\": false, \"mode\": \"btp\"}",
    "GET http://localhost:8000/api/sap/fetch   →  HTTP 503 (not configured)",
    "POST http://localhost:8000/api/upload with a real .xlsx  →  still works (Excel upload unaffected)",
]:
    bullet(step)

heading2("7.2  Frontend — Before BTP is configured")
for step in [
    "Upload screen shows drag-and-drop zone only — no 'Fetch from SAP BTP' button visible",
    "All existing tabs (Overview, Treemap, Detail Table, Compare) work normally",
]:
    bullet(step)

heading2("7.3  Backend — After BTP env vars are filled in")
for step in [
    "GET /api/sap/status  →  {\"available\": true, \"mode\": \"btp\"}",
    "GET /api/sap/fetch  →  LTVFParseResult JSON with live SAP data",
    "Check /docs (Swagger UI) to inspect response schema",
]:
    bullet(step)

heading2("7.4  Frontend — After BTP env vars are deployed to Render")
for step in [
    "Upload screen shows 'Fetch from SAP BTP' button with Live badge",
    "Clicking the button fetches live data and populates all 4 dashboard tabs identically to an Excel upload",
    "Upload history entry saved with filename 'SAP Live Data'",
    "Excel upload still works alongside the new button",
]:
    bullet(step)

info_box("Quick Test Sequence", [
    "  1.  Fill in backend/.env with BTP credentials",
    "  2.  Run:  uvicorn main:app --reload",
    "  3.  Open: http://localhost:8000/api/sap/status  →  should show available: true",
    "  4.  Open: http://localhost:8000/docs  →  test /api/sap/fetch in Swagger UI",
    "  5.  Run frontend:  cd frontend && npm run dev",
    "  6.  Open: http://localhost:3000  →  'Fetch from SAP BTP' button should appear",
    "  7.  Click the button  →  live data populates the dashboard",
], color_hex="E6F4EA", border_hex="16A34A")

doc.add_page_break()

# ── 8. FALLBACK — DIRECT SAP ──────────────────────────────────────────────────
heading1("8.  Fallback: Direct SAP Connection (Alternative Approach)")

body(
    "A direct Basic Auth approach is available as a fallback if BTP admin access "
    "or Cloud Connector setup is delayed. The code is preserved in backend/sap_client.py "
    "but is not wired into main.py in the current configuration."
)

heading2("8.1  When to use this approach")
for item in [
    "BTP admin role assignment is delayed beyond project timeline",
    "Cloud Connector installation is not feasible in the corporate network",
    "SAP system is accessible from the internet or Render.com's IP range directly",
]:
    bullet(item)

heading2("8.2  How to switch to direct mode")
body("Edit backend/main.py — change the two import lines and route logic:")
code_block([
    "# Replace BTP import with:",
    "from sap_client import fetch_ltvf_from_sap, is_sap_configured",
    "",
    "# Update /api/sap/status:",
    "return {\"available\": is_sap_configured(), \"mode\": \"direct\"}",
    "",
    "# Update /api/sap/fetch:",
    "return fetch_ltvf_from_sap()",
    "",
    "# Set these in .env instead of BTP_* vars:",
    "SAP_BASE_URL=https://<sap-host>:<port>",
    "SAP_USER=<service-account>",
    "SAP_PASSWORD=<password>",
    "SAP_CLIENT=100",
])

# ── FOOTER ────────────────────────────────────────────────────────────────────
doc.add_paragraph()
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run(
    f"LTVF Cloud Dashboard — SAP BTP Integration Document  |  "
    f"Version 1.0  |  {datetime.date.today().strftime('%d %B %Y')}  |  Business Confidential"
)
r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x99, 0x99, 0x99); r.font.italic = True

output_path = "C:/Users/I750407/ltvf-dashboard/LTVF_BTP_Integration.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
