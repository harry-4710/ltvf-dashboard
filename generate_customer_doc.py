from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

SAP_BLUE  = RGBColor(0x00, 0x33, 0x66)
MID_BLUE  = RGBColor(0x00, 0x5A, 0x9C)
DARK_GREY = RGBColor(0x33, 0x33, 0x33)
MID_GREY  = RGBColor(0x66, 0x66, 0x66)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
GREEN     = RGBColor(0x16, 0xA3, 0x4A)
AMBER     = RGBColor(0xD9, 0x77, 0x06)
RED       = RGBColor(0xDC, 0x26, 0x26)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="DDDDDD"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0"); b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(16); run.font.bold = True; run.font.color.rgb = SAP_BLUE
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1"); bot.set(qn("w:color"), "003366")
    pBdr.append(bot); pPr.append(pBdr)

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(13); run.font.bold = True; run.font.color.rgb = MID_BLUE

def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.font.size = Pt(10.5); r.font.color.rgb = DARK_GREY

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.8 + level * 0.6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(10.5); r.font.color.rgb = DARK_GREY

def numbered(text, level=0):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Cm(0.8 + level * 0.6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(10.5); r.font.color.rgb = DARK_GREY

def info_box(title, lines, color_hex="E8F0FE", border_hex="003366"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, color_hex)
    set_cell_borders(cell, border_hex)
    cell.paragraphs[0].clear()
    hr = cell.paragraphs[0].add_run(title)
    hr.font.bold = True; hr.font.size = Pt(10.5); hr.font.color.rgb = SAP_BLUE
    for line in lines:
        p = cell.add_paragraph()
        r = p.add_run(line)
        r.font.size = Pt(10.5); r.font.color.rgb = DARK_GREY
    doc.add_paragraph()

def highlight_box(title, lines, color_hex="FFF8E1", border_hex="D97706"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, color_hex)
    set_cell_borders(cell, border_hex)
    cell.paragraphs[0].clear()
    hr = cell.paragraphs[0].add_run(title)
    hr.font.bold = True; hr.font.size = Pt(10.5); hr.font.color.rgb = RGBColor(0x92, 0x40, 0x00)
    for line in lines:
        p = cell.add_paragraph()
        r = p.add_run(line)
        r.font.size = Pt(10.5); r.font.color.rgb = DARK_GREY
    doc.add_paragraph()

def green_box(title, lines):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, "E6F4EA")
    set_cell_borders(cell, "16A34A")
    cell.paragraphs[0].clear()
    hr = cell.paragraphs[0].add_run(title)
    hr.font.bold = True; hr.font.size = Pt(10.5); hr.font.color.rgb = RGBColor(0x14, 0x53, 0x2D)
    for line in lines:
        p = cell.add_paragraph()
        r = p.add_run(line)
        r.font.size = Pt(10.5); r.font.color.rgb = DARK_GREY
    doc.add_paragraph()

def data_table(headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, "003366")
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(h)
        r.font.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row in enumerate(rows):
        tr = tbl.add_row()
        bg = "F5F5F5" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            set_cell_borders(cell, "DDDDDD")
            cell.paragraphs[0].clear()
            r = cell.paragraphs[0].add_run(val)
            r.font.size = Pt(10); r.font.color.rgb = DARK_GREY
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()

def status_table(headers, rows, col_widths=None):
    """Table with colour-coded status column (last column)."""
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, "003366")
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(h)
        r.font.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        tr = tbl.add_row()
        bg = "F5F5F5" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            # Colour-code the status/last column
            if ci == len(row) - 1:
                if "Complete" in val:
                    set_cell_bg(cell, "E6F4EA")
                elif "Pending" in val:
                    set_cell_bg(cell, "FFF8E1")
                elif "Blocked" in val:
                    set_cell_bg(cell, "FEE2E2")
                else:
                    set_cell_bg(cell, bg)
            else:
                set_cell_bg(cell, bg)
            set_cell_borders(cell, "DDDDDD")
            cell.paragraphs[0].clear()
            r = cell.paragraphs[0].add_run(val)
            r.font.size = Pt(10); r.font.color.rgb = DARK_GREY
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()


# ════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ════════════════════════════════════════════════════════════════════════════
banner = doc.add_table(rows=1, cols=1)
banner.style = "Table Grid"
bc = banner.cell(0, 0)
set_cell_bg(bc, "003366")
bc.paragraphs[0].clear()
bc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
br1 = bc.paragraphs[0].add_run("\n\nLTVF Cloud Dashboard\n")
br1.font.size = Pt(30); br1.font.bold = True; br1.font.color.rgb = WHITE
br2 = bc.add_paragraph()
br2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = br2.add_run("Connecting to Your SAP System\nA Guide for Stakeholders & Team Leads\n\n")
r2.font.size = Pt(14); r2.font.color.rgb = RGBColor(0xB3, 0xD1, 0xFF)
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in [
    f"Prepared by: Hariprasad",
    f"Date: {datetime.date.today().strftime('%d %B %Y')}",
    "Version: 1.0  |  For Internal Use",
]:
    r = meta.add_run(line + "\n")
    r.font.size = Pt(11); r.font.color.rgb = DARK_GREY

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
#  1. WHAT IS THIS DOCUMENT?
# ════════════════════════════════════════════════════════════════════════════
heading1("1.  What Is This Document?")
body(
    "This document explains how the LTVF Cloud Dashboard will connect directly to your "
    "SAP system — what that means for the team, what it looks like for the user, "
    "what needs to happen before it goes live, and who needs to do what."
)
body(
    "It is written for project stakeholders, team leads, and the people who need to "
    "take action. Technical details are kept to a minimum."
)
doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
#  2. THE PROBLEM WE ARE SOLVING
# ════════════════════════════════════════════════════════════════════════════
heading1("2.  The Problem We Are Solving")

body(
    "Today, getting LTVF test results into the dashboard is a manual process:"
)
for step in [
    "A user logs into SAP GUI",
    "Runs the CNVLTVF3 report",
    "Exports the results as an Excel file",
    "Uploads that file into the dashboard",
]:
    bullet(step)

body("")
body(
    "This works, but it has drawbacks. The data is only as fresh as the last manual export. "
    "If someone forgets to re-export, the dashboard shows outdated results. "
    "It also adds unnecessary steps for every user who wants to view the latest data."
)

highlight_box("The Goal", [
    "  Replace the manual export with a single click.",
    "  Users press 'Fetch from SAP' and the dashboard instantly shows live, up-to-date results —",
    "  no SAP GUI login, no file export, no upload required.",
])

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
#  3. WHAT THE USER WILL SEE
# ════════════════════════════════════════════════════════════════════════════
heading1("3.  What the User Will See")

body(
    "Once the connection is live, the upload screen will show two options side by side:"
)

data_table(
    ["Option", "How it works", "Best for"],
    [
        ["Upload Excel File",
         "User exports from SAP GUI and uploads the .xlsx file manually — exactly as today.",
         "One-off analysis, offline use, or when SAP is unavailable."],
        ["Fetch from SAP  (new)",
         "User clicks one button. Dashboard fetches live data directly from SAP and loads instantly.",
         "Day-to-day use. Always shows the latest results."],
    ],
    col_widths=[1.5, 2.8, 2.4]
)

body(
    "The existing Excel upload is not removed. Both options are available at all times. "
    "The 'Fetch from SAP' button only appears after the connection has been set up — "
    "until then, the dashboard looks and works exactly as it does today."
)

green_box("User Experience Summary", [
    "  Before:  Open SAP GUI  →  Run report  →  Export Excel  →  Upload to dashboard  (4 steps)",
    "  After:   Open dashboard  →  Click 'Fetch from SAP'  (1 step)",
])

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
#  4. HOW THE CONNECTION WORKS
# ════════════════════════════════════════════════════════════════════════════
heading1("4.  How the Connection Works")

body(
    "The dashboard connects to SAP through SAP's own secure cloud platform — "
    "SAP Business Technology Platform (BTP). This is the standard, approved way "
    "for web applications to communicate with on-premise SAP systems."
)

heading2("4.1  The Simple Picture")

info_box("Connection Flow (Non-Technical)", [
    "",
    "   User clicks 'Fetch from SAP'",
    "          ↓",
    "   Dashboard (web app) sends a secure request to SAP BTP",
    "          ↓",
    "   SAP BTP verifies the request is authorised",
    "          ↓",
    "   SAP BTP forwards the request through a secure tunnel to the on-premise SAP system",
    "          ↓",
    "   SAP returns the CNVLTVF3 test results",
    "          ↓",
    "   Dashboard displays the live results instantly",
    "",
])

heading2("4.2  Why This Approach Is Safe")
for item in [
    "The SAP system is never directly exposed to the internet. All traffic goes through SAP's own secure tunnel.",
    "A dedicated read-only service account is used — the dashboard can only read data, never modify it.",
    "All communication is encrypted. No data travels in plain text.",
    "Credentials are stored securely on the server — they are never visible to users in the browser.",
    "This is the standard integration method recommended by SAP for connecting web applications to on-premise systems.",
]:
    bullet(item)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
#  5. WHAT NEEDS TO HAPPEN BEFORE IT GOES LIVE
# ════════════════════════════════════════════════════════════════════════════
heading1("5.  What Needs to Happen Before It Goes Live")

body(
    "The dashboard code is fully built and ready. The remaining steps are configuration "
    "and access — no further development is needed. Three teams need to take action."
)

heading2("5.1  Action Checklist")

status_table(
    ["#", "Action Required", "Who", "Status"],
    [
        ["1", "Grant Hariprasad admin access to the BTP subaccount (BDV space) so service instances and destinations can be created.", "BTP Global Account Admin", "Pending"],
        ["2", "Install the SAP Cloud Connector application on a server inside the corporate network and connect it to the BTP subaccount.", "BASIS / Infrastructure Team", "Pending"],
        ["3", "Expose the on-premise SAP system through the Cloud Connector and add the CNVLTVF3 report path to the allowed list.", "BASIS / Infrastructure Team", "Pending"],
        ["4", "Create a read-only SAP service account with access to the CNVLTVF3 report. Share credentials securely with Hariprasad.", "BASIS Team", "Pending"],
        ["5", "Create the required BTP service instances (Destination and Connectivity) and set up a named connection to the SAP system.", "Hariprasad  (after steps 1–4)", "Pending"],
        ["6", "Enter the connection credentials into the dashboard backend and redeploy. Test end-to-end.", "Hariprasad  (after step 5)", "Pending"],
    ],
    col_widths=[0.3, 3.5, 1.8, 1.1]
)

highlight_box("Currently Blocked On", [
    "  Steps 1, 2, 3, and 4 require action from the BTP Admin and BASIS team before anything else can proceed.",
    "  A request email has been sent covering all three items.",
])

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
#  6. WHO NEEDS TO DO WHAT
# ════════════════════════════════════════════════════════════════════════════
heading1("6.  Who Needs to Do What")

heading2("6.1  BTP Global Account Administrator")
body("One action required:")
numbered("Assign the following roles to user Hariprasad (HV) in the BDV space subaccount:")
bullet("Subaccount Administrator", level=1)
bullet("Connectivity and Destination Administrator", level=1)
body(
    "This gives Hariprasad the ability to create the connection configuration in BTP. "
    "It does not grant any access to the SAP system itself."
)

heading2("6.2  BASIS / Infrastructure Team")
body("Two actions required:")
numbered("Install the SAP Cloud Connector on a server inside the corporate network and link it to the BTP subaccount BDV space. This is a one-time setup.")
numbered("Add the on-premise SAP system to the Cloud Connector and permit access to the CNVLTVF3 report path. No inbound firewall changes are needed — the connector uses an outbound-only tunnel.")

body(
    "The Cloud Connector is a standard SAP tool, freely available from SAP, and widely "
    "used across SAP BTP projects. Installation typically takes 1–2 hours."
)

heading2("6.3  BASIS Team (Service Account)")
body("One action required:")
numbered(
    "Create a read-only technical user in the SAP system with access to the CNVLTVF3 report and its data service. "
    "Share the username, password, SAP system URL, and client number securely with Hariprasad."
)
body(
    "This account will be used only by the dashboard for automated data reads. "
    "It should have no write permissions."
)

heading2("6.4  Hariprasad (after above steps are complete)")
body("Three actions to complete once access is granted:")
numbered("Create the BTP connection configuration (Destination and Connectivity service instances) in the BTP cockpit — approximately 30 minutes.")
numbered("Enter the connection details into the dashboard and redeploy.")
numbered("Run end-to-end testing and confirm the 'Fetch from SAP' button works correctly.")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
#  7. TIMELINE ESTIMATE
# ════════════════════════════════════════════════════════════════════════════
heading1("7.  Estimated Timeline")

body(
    "Once all access and configuration steps are completed, the 'Fetch from SAP' feature "
    "can be live within the same day. The timeline below assumes normal response times."
)

data_table(
    ["Step", "Estimated Time", "Depends On"],
    [
        ["BTP admin roles assigned",              "Same day as request",     "BTP Admin availability"],
        ["Cloud Connector installed & connected", "1–2 days",                "BASIS team availability"],
        ["SAP service account created",           "1–2 days",                "BASIS team availability"],
        ["BTP configuration by Hariprasad",       "2–3 hours",               "Steps above complete"],
        ["Testing & go-live",                     "2–3 hours",               "BTP configuration done"],
        ["Total (from approvals to go-live)",     "2–4 business days",       "All teams responding"],
    ],
    col_widths=[2.5, 1.8, 2.4]
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
#  8. FREQUENTLY ASKED QUESTIONS
# ════════════════════════════════════════════════════════════════════════════
heading1("8.  Frequently Asked Questions")

heading2("Will this affect the current way of using the dashboard?")
body(
    "No. The Excel upload option remains exactly as it is today. The 'Fetch from SAP' button "
    "is an additional option. Users can continue uploading files if they prefer."
)

heading2("Is the SAP data safe when it travels through this connection?")
body(
    "Yes. The connection uses SAP's own secure platform (BTP) with encrypted communication. "
    "SAP data never travels over the open internet unprotected. The approach is the same "
    "method SAP recommends for all web-to-SAP integrations."
)

heading2("Can the dashboard change or delete any SAP data?")
body(
    "No. The service account used for the connection is strictly read-only. "
    "The dashboard can only retrieve data — it has no ability to create, modify, or delete "
    "anything in the SAP system."
)

heading2("What happens if the SAP system is down or unreachable?")
body(
    "The 'Fetch from SAP' button will show an error message. The Excel upload option "
    "continues to work normally regardless of SAP availability."
)

heading2("Does this require any changes to SAP itself?")
body(
    "No changes to SAP are needed. The only requirement is a read-only service account "
    "and a network path from the Cloud Connector server to the SAP system — both of "
    "which already exist in most corporate environments."
)

heading2("Who should be contacted if something goes wrong after go-live?")
body(
    "Hariprasad is the point of contact for the dashboard. For SAP system access issues, "
    "the BASIS team should be involved. For BTP platform issues, the BTP Global Account "
    "Administrator can assist."
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
#  9. SUMMARY
# ════════════════════════════════════════════════════════════════════════════
heading1("9.  Summary")

green_box("What Is Ready", [
    "  ✔  Dashboard is fully built — live SAP connection is coded and tested in development",
    "  ✔  Excel upload continues to work as normal — no disruption to current users",
    "  ✔  Security design reviewed — read-only access, encrypted connection, no exposed credentials",
    "  ✔  Request email sent to BTP Admin and BASIS team",
])

highlight_box("What Is Needed to Go Live", [
    "  1.  BTP Admin → assign roles to Hariprasad in BDV space subaccount",
    "  2.  BASIS / Infrastructure → install and configure SAP Cloud Connector",
    "  3.  BASIS → create read-only SAP service account for the dashboard",
    "  Once the above are done → Hariprasad completes configuration and testing within hours.",
])

body(
    "The goal is to give every team member instant access to live CNVLTVF3 results "
    "without needing SAP GUI access or manual file exports — making the migration "
    "quality review process faster and more reliable for everyone."
)

# ── FOOTER ───────────────────────────────────────────────────────────────────
doc.add_paragraph()
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run(
    f"LTVF Cloud Dashboard — SAP Connection Guide  |  "
    f"Version 1.0  |  {datetime.date.today().strftime('%d %B %Y')}  |  Internal Use Only"
)
r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x99, 0x99, 0x99); r.font.italic = True

output_path = "C:/Users/I750407/ltvf-dashboard/LTVF_SAP_Connection_Guide.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
