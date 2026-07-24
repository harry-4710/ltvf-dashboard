from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Colour palette ────────────────────────────────────────────────────────────
SAP_BLUE    = RGBColor(0x00, 0x33, 0x66)
MID_BLUE    = RGBColor(0x00, 0x5A, 0x9C)
ACCENT_BLUE = RGBColor(0x0A, 0x84, 0xFF)
LIGHT_GREY  = RGBColor(0xF5, 0xF5, 0xF5)
DARK_GREY   = RGBColor(0x33, 0x33, 0x33)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
GREEN       = RGBColor(0x16, 0xA3, 0x4A)
AMBER       = RGBColor(0xD9, 0x77, 0x06)
RED         = RGBColor(0xDC, 0x26, 0x26)

# ── Helper utilities ──────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color="DDDDDD"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),   "single")
        border.set(qn("w:sz"),    "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def heading1(text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.size  = Pt(16)
    run.font.bold  = True
    run.font.color.rgb = SAP_BLUE
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "003366")
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


def heading2(text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.size  = Pt(13)
    run.font.bold  = True
    run.font.color.rgb = MID_BLUE
    return p


def heading3(text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size  = Pt(11)
    run.font.bold  = True
    run.font.color.rgb = DARK_GREY
    return p


def body(text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = DARK_GREY
    return p


def bullet(text: str, level: int = 0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.8 + level * 0.6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = DARK_GREY
    return p


def code_block(lines: list):
    p = doc.add_paragraph()
    r = p.add_run("\n".join(lines))
    r.font.name = "Courier New"
    r.font.size = Pt(8.5)
    r.font.color.rgb = DARK_GREY
    doc.add_paragraph()


def info_box(title: str, lines: list, color_hex="E8F0FE", border_hex="003366"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, color_hex)
    set_cell_borders(cell, border_hex)
    cell.paragraphs[0].clear()
    hrun = cell.paragraphs[0].add_run(title)
    hrun.font.bold = True
    hrun.font.size = Pt(10.5)
    hrun.font.color.rgb = SAP_BLUE
    for line in lines:
        p = cell.add_paragraph()
        r = p.add_run(line)
        r.font.size = Pt(10)
        r.font.color.rgb = DARK_GREY
    doc.add_paragraph()


def data_table(headers: list, rows: list, col_widths: list = None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, "003366")
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(h)
        r.font.bold  = True
        r.font.color.rgb = WHITE
        r.font.size  = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row in enumerate(rows):
        tr = tbl.add_row()
        bg = "F5F5F5" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            set_cell_borders(cell, "DDDDDD")
            cell.paragraphs[0].clear()
            r = cell.paragraphs[0].add_run(val)
            r.font.size = Pt(10)
            r.font.color.rgb = DARK_GREY
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════

banner = doc.add_table(rows=1, cols=1)
banner.style = "Table Grid"
bc = banner.cell(0, 0)
set_cell_bg(bc, "003366")
bc.paragraphs[0].clear()
bc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
br1 = bc.paragraphs[0].add_run("\n\nLTVF Cloud Dashboard\n")
br1.font.size  = Pt(28)
br1.font.bold  = True
br1.font.color.rgb = WHITE
br2 = bc.add_paragraph()
br2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2  = br2.add_run("Technical Architecture & Implementation Reference\n\n")
r2.font.size  = Pt(14)
r2.font.color.rgb = RGBColor(0xB3, 0xD1, 0xFF)
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in [
    "SAP CNVLTVF3 Migration Test Results Dashboard",
    f"Date: {datetime.date.today().strftime('%d %B %Y')}",
    "Version: 2.0",
    "Classification: Business Confidential",
]:
    r = meta.add_run(line + "\n")
    r.font.size = Pt(11)
    r.font.color.rgb = DARK_GREY

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════════════════════
heading1("Table of Contents")
toc_items = [
    "1.  Executive Summary",
    "2.  Project Overview & Purpose",
    "3.  Solution Architecture",
    "4.  Frontend Architecture (React + TypeScript)",
    "5.  Backend Architecture (Python FastAPI)",
    "6.  Excel Parser — Core Logic",
    "7.  Data Models",
    "8.  API Reference",
    "9.  Deployment Architecture",
    "10. Security Design",
    "11. Configuration & Environment Variables",
    "12. Local Development Setup",
    "13. Production Deployment",
    "14. Glossary",
]
for item in toc_items:
    bullet(item)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════
heading1("1.  Executive Summary")

body(
    "The LTVF Cloud Dashboard is a full-stack web application built to visualise the output "
    "of SAP transaction CNVLTVF3 — the Long-Term Value Forecast test framework used during "
    "SAP S/4HANA system conversions. Users upload the Excel export produced by CNVLTVF3 and "
    "immediately receive an interactive, browser-based dashboard showing test case match rates, "
    "pass/warn/fail breakdowns, hierarchical treemaps, and a filterable detail table."
)
body(
    "The application is stateless — no database is used. All data is parsed on the fly from "
    "the uploaded file, returned to the browser as JSON, and optionally cached in browser "
    "localStorage for quick reload. The frontend is deployed on Vercel; the backend on "
    "Render.com. Local development uses Docker Compose."
)

info_box(
    "Key Facts",
    [
        "  •  Frontend:  React 18 + TypeScript + Vite + Tailwind CSS + Recharts + AG Grid",
        "  •  Backend:   Python 3.12 + FastAPI 0.111 + pandas + openpyxl",
        "  •  Database:  None — stateless, file-upload-only model",
        "  •  Deployment: Vercel (frontend)  +  Render.com (backend)",
        "  •  Local dev:  Docker Compose (frontend :3000, backend :8000)",
        "  •  Source:     C:\\Users\\I750407\\ltvf-dashboard",
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  2. PROJECT OVERVIEW & PURPOSE
# ═══════════════════════════════════════════════════════════════════════════════
heading1("2.  Project Overview & Purpose")

heading2("2.1  Business Context")
body(
    "During an SAP system conversion (e.g. ECC to S/4HANA), the CNVLTVF3 test framework "
    "runs thousands of test cases that compare source and target data records. Each test case "
    "produces a match rate (%) — how closely the target data aligns with the source. The "
    "raw output is an Excel file with a fixed 20-column layout and a hierarchical row structure "
    "using '>> path > subpath >' prefixes for group rows."
)
body(
    "This dashboard replaces the need to manually read the Excel file. It parses the file "
    "server-side and renders the results as an interactive, colour-coded web application "
    "accessible from any browser without SAP GUI."
)

heading2("2.2  Thresholds & Traffic Light System")
data_table(
    ["Status", "Condition", "Colour"],
    [
        ["Pass",    "match rate >= 95%  (default)", "Green"],
        ["Warn",    "80% <= match rate < 95%",       "Amber"],
        ["Fail",    "match rate < 80%",               "Red"],
    ],
    col_widths=[1.2, 2.8, 1.2]
)
body("Both thresholds are configurable at runtime via the Threshold Panel in the dashboard header — no code change required.")

heading2("2.3  Repository Structure")
code_block([
    "ltvf-dashboard/",
    "├── frontend/                 React + TypeScript SPA (Vite)",
    "│   ├── src/",
    "│   │   ├── App.tsx           Root component — tab state, upload logic, theme",
    "│   │   ├── api/",
    "│   │   │   └── sapApi.ts     uploadLTVF() — posts file to /api/upload",
    "│   │   ├── components/       All UI components (12 files)",
    "│   │   ├── hooks/",
    "│   │   │   └── useCountUp.ts Animated number count-up hook",
    "│   │   ├── types/",
    "│   │   │   └── ltvf.ts       TypeScript interfaces mirroring Pydantic models",
    "│   │   └── utils/",
    "│   │       └── history.ts    localStorage upload history (max 5 entries)",
    "│   ├── vite.config.ts        Dev server port 3000, /api proxy to :8000",
    "│   ├── tailwind.config.js    darkMode: 'class'",
    "│   ├── vercel.json           Vercel deployment + SPA rewrites",
    "│   ├── Dockerfile            Multi-stage: node build → nginx:alpine serve",
    "│   └── nginx.conf            SPA fallback + /api/ reverse proxy to backend",
    "├── backend/",
    "│   ├── main.py               FastAPI app — CORS + 2 routes",
    "│   ├── excel_parser.py       Core parser — Pydantic models + parse_excel()",
    "│   ├── models.py             Legacy Pydantic models (OData era, unused)",
    "│   ├── mock_data.py          Legacy mock data (unused)",
    "│   ├── sap_client.py         Legacy SAP OData client (not wired, reserved)",
    "│   ├── requirements.txt      Python dependencies",
    "│   ├── Dockerfile            python:3.12-slim + uvicorn",
    "│   ├── render.yaml           Render.com deployment config",
    "│   └── .env.example          Environment variable template",
    "├── docker-compose.yml        Local orchestration (backend :8000, frontend :3000)",
    "├── generate_doc.py           This document generator",
    "├── LTVF_Cloud_Dashboard_Architecture.docx   Generated output",
    "└── .gitignore",
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  3. SOLUTION ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════════
heading1("3.  Solution Architecture")

heading2("3.1  Architecture Diagram")
code_block([
    "┌─────────────────────────────────────────────────────────────────────┐",
    "│                    USER BROWSER (any device)                        │",
    "│                                                                     │",
    "│   React SPA  ──  Tab: Overview | Treemap | Detail Table | Compare  │",
    "│   Recharts  ·  AG Grid  ·  Tailwind CSS  ·  Dark / Light mode      │",
    "└────────────────────────────┬────────────────────────────────────────┘",
    "                             │  HTTPS  POST /api/upload  (.xlsx/.xls) ",
    "                             │  HTTPS  GET  /api/health               ",
    "┌────────────────────────────▼────────────────────────────────────────┐",
    "│                 FASTAPI BACKEND  (Render.com)                       │",
    "│                                                                     │",
    "│   main.py          ──  CORS middleware + route definitions          │",
    "│   excel_parser.py  ──  pandas parse → hierarchy build → Pydantic   │",
    "│                                                                     │",
    "│   Stateless — no database, no session, no file storage             │",
    "└─────────────────────────────────────────────────────────────────────┘",
    "",
    "  Data flow:",
    "  1. User drags .xlsx onto UploadZone (or clicks to browse)",
    "  2. Browser POSTs file to POST /api/upload via axios",
    "  3. FastAPI reads file bytes → passes to parse_excel()",
    "  4. parse_excel() builds hierarchy, classifies pass/warn/fail",
    "  5. Returns LTVFParseResult JSON to browser",
    "  6. React renders charts, cards, table from the JSON",
    "  7. Result cached in localStorage for history (max 5 entries)",
])

heading2("3.2  Key Design Decisions")
data_table(
    ["Decision", "Choice", "Rationale"],
    [
        ["No database",          "Stateless file upload",        "CNVLTVF3 results are point-in-time snapshots; no persistence needed. Eliminates DB ops overhead."],
        ["No URL router",        "Tab state in App.tsx",         "Single-page with 4 tabs; no bookmarkable views needed. Keeps bundle small."],
        ["localStorage history", "Up to 5 recent uploads",       "Allows quick reload of recent files without re-uploading, survives page refresh."],
        ["Configurable thresholds", "Runtime sliders",           "Pass/warn cutoffs vary by project; no redeploy needed to adjust."],
        ["Dark mode",            "Tailwind class strategy",      "User preference persisted in localStorage; respects OS prefers-color-scheme on first load."],
        ["Excel column mapping", "Position-based (not by name)", "CNVLTVF3 export has no guaranteed header names; mapping by column index is robust."],
    ],
    col_widths=[1.6, 1.6, 3.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  4. FRONTEND ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════════
heading1("4.  Frontend Architecture")

heading2("4.1  Technology Stack")
data_table(
    ["Package", "Version", "Role"],
    [
        ["react",              "18.3.1",  "UI library — component rendering"],
        ["react-dom",          "18.3.1",  "DOM renderer"],
        ["typescript",         "5.4.5",   "Static typing across all source files"],
        ["vite",               "5.3.1",   "Dev server (port 3000) + production bundler"],
        ["tailwindcss",        "3.4.4",   "Utility-first CSS, dark mode via class strategy"],
        ["recharts",           "2.12.7",  "Charts: PieChart, BarChart, Treemap"],
        ["ag-grid-react",      "31.3.2",  "Detail table with search, column toggle, CSV export"],
        ["ag-grid-community",  "31.3.2",  "AG Grid core (community edition)"],
        ["axios",              "1.7.2",   "HTTP client for /api/upload calls"],
        ["lucide-react",       "0.395.0", "Icons: Database, Upload, Sun, Moon, GitCompare, Clock, Printer"],
        ["autoprefixer",       "10.4.19", "PostCSS plugin for Tailwind"],
        ["@vitejs/plugin-react","4.3.1",  "Vite React/JSX transform plugin"],
    ],
    col_widths=[2.0, 1.0, 3.7]
)

heading2("4.2  Tab Navigation")
body(
    "Navigation is managed entirely by a tab state variable in App.tsx — there is no URL "
    "router. Tab buttons are rendered in the sticky header bar once a file is loaded."
)
data_table(
    ["Tab key", "Label", "Content"],
    [
        ["overview",  "Overview",     "SummaryCards + RateDonut + SectionChart + VolumeChart + FailChart"],
        ["treemap",   "Treemap",      "TreemapChart — Recharts Treemap colored by pass/warn/fail rate"],
        ["table",     "Detail Table", "LTVFTable — AG Grid with search, column visibility toggle, CSV export"],
        ["compare",   "Compare",      "ComparePanel — side-by-side two-file comparison with delta table"],
    ],
    col_widths=[1.2, 1.2, 4.3]
)

heading2("4.3  Component Catalogue")
data_table(
    ["Component file", "Purpose"],
    [
        ["UploadZone.tsx",     "Landing screen — drag-and-drop or click-to-browse for .xlsx/.xls files"],
        ["SummaryCards.tsx",   "8 animated KPI cards: Overall Rate, Equal, Differences, Missing, Unexpected, Source Volume, Target Volume, Test Cases"],
        ["RateDonut.tsx",      "Recharts PieChart donut — pass/warn/fail test case counts with overall % in centre"],
        ["SectionChart.tsx",   "Horizontal Recharts BarChart — match rate per section; clickable to set section filter"],
        ["VolumeChart.tsx",    "Grouped bar chart — Equal / Diff / Missing / Unexpected record counts"],
        ["FailChart.tsx",      "Horizontal bar chart — 15 lowest-performing test cases"],
        ["LTVFTable.tsx",      "AG Grid table — hierarchical rows, quick-search, column visibility toggle, CSV export"],
        ["TreemapChart.tsx",   "Recharts Treemap with custom SVG content renderer, coloured by match rate"],
        ["ComparePanel.tsx",   "Two-file comparison: delta table + side-by-side donuts + side-by-side summary cards"],
        ["ThresholdPanel.tsx", "Popover with range sliders to configure Pass (default 95%) and Warn (default 80%) cutoffs"],
        ["FilterChips.tsx",    "Section-level filter chips shown above all chart tabs — click to filter all charts"],
        ["UploadHistory.tsx",  "Dropdown panel listing up to 5 recent uploads from localStorage"],
    ],
    col_widths=[2.1, 4.6]
)

heading2("4.4  Key Utilities & Hooks")
data_table(
    ["File", "Purpose"],
    [
        ["src/api/sapApi.ts",      "uploadLTVF(file: File): Promise<LTVFParseResult> — POSTs multipart form to /api/upload. In dev, Vite proxy rewrites /api to http://localhost:8000. In production reads VITE_API_URL env var."],
        ["src/types/ltvf.ts",      "TypeScript interfaces LTVFRow, LTVFSummary, LTVFParseResult — mirror the backend Pydantic models exactly."],
        ["src/utils/history.ts",   "localStorage-backed upload history (max 5). Functions: loadHistory(), saveToHistory(), deleteFromHistory(). Each entry stores: filename, timestamp, systemTag, full LTVFParseResult."],
        ["src/hooks/useCountUp.ts","Custom hook for animated number count-up using requestAnimationFrame with ease-out cubic easing (900ms duration)."],
    ],
    col_widths=[2.1, 4.6]
)

heading2("4.5  App-level State (App.tsx)")
data_table(
    ["State variable", "Type", "Purpose"],
    [
        ["data",            "LTVFParseResult | null",  "Current uploaded file parse result"],
        ["compareData",     "LTVFParseResult | null",  "Second file for Compare tab"],
        ["loading",         "boolean",                 "Primary upload in progress"],
        ["compareLoading",  "boolean",                 "Compare upload in progress"],
        ["error",           "string | null",           "Error banner message"],
        ["tab",             "Tab",                     "Active tab: overview | table | treemap | compare"],
        ["selectedSection", "string | null",           "Active section filter — applied to all charts"],
        ["thresholds",      "{ pass: 95, warn: 80 }",  "Pass/warn cutoff percentages"],
        ["systemTag",       "string",                  "Optional label (e.g. QSL, PRD) shown in header"],
        ["history",         "HistoryEntry[]",          "Up to 5 recent uploads from localStorage"],
        ["uploadedAt",      "Date | null",             "Timestamp of current upload shown in header"],
        ["dark",            "boolean",                 "Theme state — persisted in localStorage"],
    ],
    col_widths=[1.9, 2.1, 2.7]
)

heading2("4.6  Theme System")
body(
    "Dark mode is implemented via Tailwind's 'class' strategy. The root <div> receives the "
    "class 'dark' when dark mode is active. The useTheme hook reads/writes 'ltvf-theme' in "
    "localStorage and falls back to window.matchMedia('(prefers-color-scheme: dark)') on "
    "first load. All components receive a 'dark: boolean' prop and apply conditional Tailwind "
    "classes accordingly."
)

heading2("4.7  Print / PDF Export")
body(
    "A Printer icon in the header calls window.print(). A 'no-print' CSS class is applied to "
    "the sticky header and error banner so they are excluded from the print layout, leaving "
    "only the dashboard content."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  5. BACKEND ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════════
heading1("5.  Backend Architecture")

heading2("5.1  Technology Stack")
data_table(
    ["Package", "Version", "Role"],
    [
        ["fastapi",            "0.111.0", "Web framework — routing, dependency injection, OpenAPI docs"],
        ["uvicorn[standard]",  "0.30.0",  "ASGI server (production: guarded by Render.com proxy)"],
        ["pydantic",           "2.7.1",   "Data validation and serialisation for request/response models"],
        ["pandas",             "2.2.2",   "Excel file reading and dataframe manipulation"],
        ["openpyxl",           "3.1.4",   "pandas Excel engine for .xlsx files"],
        ["python-multipart",   "0.0.9",   "Multipart form data support (required for UploadFile)"],
        ["python-dotenv",      "1.0.1",   "Loads .env file into os.environ at startup"],
        ["httpx",              "0.27.0",  "Async HTTP client (reserved for SAP OData client)"],
        ["requests",           "2.32.3",  "Sync HTTP client (legacy sap_client.py)"],
        ["zeep",               "4.2.1",   "SOAP client (legacy sap_client.py, not wired to main)"],
    ],
    col_widths=[2.0, 1.0, 3.7]
)

heading2("5.2  Application Entry Point (main.py)")
body(
    "main.py creates the FastAPI application, configures CORS middleware, and declares two "
    "routes. It imports parse_excel and LTVFParseResult from excel_parser.py — these are the "
    "only runtime dependencies."
)
code_block([
    "# CORS configuration",
    "_origins_env = os.getenv('ALLOWED_ORIGINS', '*')",
    "_origins = [o.strip() for o in _origins_env.split(',')]",
    "  if _origins_env != '*' else ['*']",
    "",
    "# Routes",
    "POST  /api/upload  — accepts UploadFile (.xlsx/.xls), calls parse_excel(), returns LTVFParseResult",
    "GET   /api/health  — returns {'status': 'ok', 'version': '2.0.0'}",
])

heading2("5.3  CORS Policy")
body(
    "In local development ALLOWED_ORIGINS defaults to '*' (wildcard). In production on "
    "Render.com, set ALLOWED_ORIGINS to the Vercel frontend URL "
    "(e.g. https://ltvf-dashboard.vercel.app) to restrict API access to the known origin only."
)

heading2("5.4  Error Handling")
body("The /api/upload endpoint has two explicit error paths:")
bullet("HTTP 400 — file extension is not .xlsx or .xls")
bullet("HTTP 422 — parse_excel() raises any exception; full traceback included in the detail field for debugging")

heading2("5.5  Legacy Modules (not wired to main.py)")
data_table(
    ["File", "Status", "Description"],
    [
        ["sap_client.py",  "Preserved, not used", "SAP OData client — calls /sap/opu/odata/sap/CNVLTVF3_SRV/LTVFResultSet with HTTP Basic Auth. Was the original data source before the Excel-upload model was adopted."],
        ["models.py",      "Preserved, not used", "Older Pydantic models from the OData era (LTVFItem with gl_account, cost_center, delta fields)."],
        ["mock_data.py",   "Preserved, not used", "Static mock dataset with SAP GL accounts, cost centres, business partners — used during early development."],
    ],
    col_widths=[1.4, 1.6, 3.7]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  6. EXCEL PARSER — CORE LOGIC
# ═══════════════════════════════════════════════════════════════════════════════
heading1("6.  Excel Parser — Core Logic")

heading2("6.1  Input Format")
body(
    "The parser expects the Excel file produced by SAP transaction CNVLTVF3. The file has "
    "a fixed 20-column layout. Column names are not reliable — the parser maps by position:"
)
data_table(
    ["Index", "Internal name", "Description"],
    [
        ["0",  "test_name",   "Test case or group name. Groups start with '>>'"],
        ["1",  "rate_pct",    "Match rate percentage (0–100)"],
        ["2",  "eval",        "Evaluation flag (unused in current UI)"],
        ["3",  "diff",        "Count of differing records"],
        ["4",  "accept",      "Count of accepted differences"],
        ["5",  "missing",     "Records in source but not in target"],
        ["6",  "unexpected",  "Records in target but not in source"],
        ["7",  "equal",       "Records that match exactly"],
        ["8",  "oos_src",     "Out-of-scope — source"],
        ["9",  "oos_trg",     "Out-of-scope — target"],
        ["10", "_local_lbl",  "Label column (skipped)"],
        ["11", "local",       "Local record count"],
        ["12", "_src_lbl",    "Label column (skipped)"],
        ["13", "source",      "Source volume"],
        ["14", "_src1_lbl",   "Label column (skipped)"],
        ["15", "source1",     "Source1 volume"],
        ["16", "_src2_lbl",   "Label column (skipped)"],
        ["17", "source2",     "Source2 volume"],
        ["18", "_tgt_lbl",    "Label column (skipped)"],
        ["19", "target",      "Target volume"],
    ],
    col_widths=[0.6, 1.3, 4.8]
)

heading2("6.2  Parsing Algorithm")
body("The parse_excel() function follows these steps in order:")

for step in [
    "1.  Read file bytes into a pandas DataFrame using header=0 (first row as header).",
    "2.  Rename columns by position using the 20-element _col_names list. Columns beyond index 19 are ignored.",
    "3.  Ensure all expected numeric columns exist; fill missing ones with None.",
    "4.  Extract row 0 as the grand-total summary row (overall rate, totals for equal/diff/missing/unexpected/source/target). Drop row 0 from the DataFrame.",
    "5.  Drop any rows where test_name is NaN.",
    "6.  Iterate remaining rows. For each row:",
    "      a.  If test_name starts with '>>' → it is a group/hierarchy row.",
    "          Parse path segments by splitting on '>': e.g. '>> 02. Test Cases > 1 MASTER DATA > 1.2 FI-GL >'",
    "          Trim the path_stack to match the depth of the new group.",
    "          Assign parent_id from top of path_stack. level = len(path_parts) - 1.",
    "          Push this group onto path_stack.",
    "          Collect top-level path segments into the sections list.",
    "      b.  If not a group → leaf test case row.",
    "          parent_id from top of path_stack. level = len(path_stack).",
    "          Classify: rate >= 95 → pass, >= 80 → warn, else fail.",
    "7.  Return LTVFParseResult with summary + flat rows list + sections list.",
]:
    bullet(step if not step.startswith("     ") else step.strip(), level=1 if step.startswith("     ") else 0)

heading2("6.3  Hierarchy Example")
code_block([
    "Excel rows (raw test_name column):",
    "",
    ">> 02. Test Cases > 1 MASTER DATA >            ← group, level 0, id=row_1, parent=None",
    ">> 02. Test Cases > 1 MASTER DATA > 1.1 FI-GL > ← group, level 1, id=row_2, parent=row_1",
    "GL Account 100000                               ← leaf,  level 2, id=row_3, parent=row_2",
    "GL Account 100010                               ← leaf,  level 2, id=row_4, parent=row_2",
    ">> 02. Test Cases > 1 MASTER DATA > 1.2 CO >   ← group, level 1, id=row_5, parent=row_1",
    "Cost Centre 1000                                ← leaf,  level 2, id=row_6, parent=row_5",
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  7. DATA MODELS
# ═══════════════════════════════════════════════════════════════════════════════
heading1("7.  Data Models")

heading2("7.1  LTVFRow")
body("Represents a single row from the parsed Excel file — either a group header or a leaf test case.")
data_table(
    ["Field", "Type", "Description"],
    [
        ["id",          "str",            "Unique row identifier (e.g. 'row_42')"],
        ["parent_id",   "str | None",     "ID of the parent group row; None for top-level groups"],
        ["level",       "int",            "Depth in the hierarchy (0 = top-level group)"],
        ["test_name",   "str",            "Display name of the test case or group"],
        ["full_path",   "str",            "Full '>' separated path (e.g. '02. Test Cases > 1 MASTER DATA > GL Account 100000')"],
        ["is_group",    "bool",           "True for group/hierarchy rows (>> prefix); False for leaf test cases"],
        ["rate_pct",    "float | None",   "Match rate percentage"],
        ["diff",        "int | None",     "Count of differing records"],
        ["accept",      "int | None",     "Count of accepted differences"],
        ["missing",     "int | None",     "Records in source but not in target"],
        ["unexpected",  "int | None",     "Records in target but not in source"],
        ["equal",       "int | None",     "Records matching exactly"],
        ["oos_src",     "int | None",     "Out-of-scope source records"],
        ["oos_trg",     "int | None",     "Out-of-scope target records"],
        ["local",       "int | None",     "Local record count"],
        ["source",      "int | None",     "Source volume"],
        ["source1",     "int | None",     "Source1 volume"],
        ["source2",     "int | None",     "Source2 volume"],
        ["target",      "int | None",     "Target volume"],
    ],
    col_widths=[1.4, 1.2, 4.1]
)

heading2("7.2  LTVFSummary")
body("Grand-total row extracted from the first row of the Excel file.")
data_table(
    ["Field", "Type", "Description"],
    [
        ["overall_rate",     "float",  "Overall match rate from grand-total row"],
        ["total_equal",      "int",    "Total equal records across all test cases"],
        ["total_diff",       "int",    "Total differing records"],
        ["total_missing",    "int",    "Total missing records"],
        ["total_unexpected", "int",    "Total unexpected records"],
        ["total_source",     "int",    "Total source volume"],
        ["total_target",     "int",    "Total target volume"],
        ["total_rows",       "int",    "Count of leaf test cases (pass + warn + fail)"],
        ["pass_count",       "int",    "Test cases with rate >= 95%"],
        ["warn_count",       "int",    "Test cases with 80% <= rate < 95%"],
        ["fail_count",       "int",    "Test cases with rate < 80%"],
    ],
    col_widths=[1.8, 0.8, 4.1]
)

heading2("7.3  LTVFParseResult  (API response root)")
data_table(
    ["Field", "Type", "Description"],
    [
        ["filename",  "str",           "Original uploaded filename"],
        ["summary",   "LTVFSummary",   "Grand-total summary object"],
        ["rows",      "list[LTVFRow]", "Flat list of all rows in document order (groups + leaves)"],
        ["sections",  "list[str]",     "Ordered list of unique top-level path segments for filter chips"],
    ],
    col_widths=[1.2, 1.6, 3.9]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  8. API REFERENCE
# ═══════════════════════════════════════════════════════════════════════════════
heading1("8.  API Reference")

heading2("8.1  POST /api/upload")
body("Accepts a multipart file upload, parses the Excel file, and returns the structured result.")
data_table(
    ["Property", "Value"],
    [
        ["URL",             "/api/upload"],
        ["Method",          "POST"],
        ["Content-Type",    "multipart/form-data"],
        ["Field name",      "file"],
        ["Accepted types",  ".xlsx, .xls"],
        ["Response",        "200 OK — LTVFParseResult JSON"],
        ["Error 400",       "File extension not .xlsx or .xls"],
        ["Error 422",       "Excel parsing failed — detail includes traceback"],
    ],
    col_widths=[1.8, 4.9]
)

heading2("8.2  GET /api/health")
body("Health check endpoint used by Render.com and monitoring tools.")
data_table(
    ["Property", "Value"],
    [
        ["URL",      "/api/health"],
        ["Method",   "GET"],
        ["Response", '200 OK — {"status": "ok", "version": "2.0.0"}'],
    ],
    col_widths=[1.8, 4.9]
)

heading2("8.3  Auto-generated Docs")
body(
    "FastAPI automatically generates interactive API documentation at /docs (Swagger UI) "
    "and /redoc (ReDoc). Available in both development and production unless explicitly disabled."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  9. DEPLOYMENT ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════════
heading1("9.  Deployment Architecture")

heading2("9.1  Docker Compose (Local Development)")
body("docker-compose.yml orchestrates both services locally:")
data_table(
    ["Service", "Build context", "Port", "Notes"],
    [
        ["backend",   "./backend",   "8000:8000",  "Loads ./backend/.env; sets USE_MOCK=true by default"],
        ["frontend",  "./frontend",  "3000:80",    "Multi-stage build: node:20-alpine builds Vite bundle → nginx:alpine serves it"],
    ],
    col_widths=[1.1, 1.5, 1.3, 2.8]
)
code_block([
    "# Start local dev environment",
    "docker-compose up --build",
    "",
    "# Frontend available at:  http://localhost:3000",
    "# Backend API at:         http://localhost:8000",
    "# Swagger docs at:        http://localhost:8000/docs",
])

heading2("9.2  Backend Dockerfile")
code_block([
    "FROM python:3.12-slim",
    "WORKDIR /app",
    "COPY requirements.txt .",
    "RUN pip install --no-cache-dir -r requirements.txt",
    "COPY . .",
    "CMD [\"uvicorn\", \"main:app\", \"--host\", \"0.0.0.0\", \"--port\", \"8000\"]",
])

heading2("9.3  Frontend Dockerfile (Multi-stage)")
code_block([
    "# Stage 1 — Build",
    "FROM node:20-alpine AS builder",
    "WORKDIR /app",
    "COPY package*.json ./",
    "RUN npm install",
    "COPY . .",
    "RUN npm run build           # produces /app/dist/",
    "",
    "# Stage 2 — Serve",
    "FROM nginx:alpine",
    "COPY --from=builder /app/dist /usr/share/nginx/html",
    "COPY nginx.conf /etc/nginx/conf.d/default.conf",
    "EXPOSE 80",
])

heading2("9.4  nginx.conf (Production)")
code_block([
    "server {",
    "  listen 80;",
    "  root /usr/share/nginx/html;",
    "  index index.html;",
    "",
    "  # SPA fallback — all unknown paths serve index.html",
    "  location / {",
    "    try_files $uri /index.html;",
    "  }",
    "",
    "  # Reverse-proxy API calls to backend container",
    "  location /api/ {",
    "    proxy_pass http://backend:8000/api/;",
    "  }",
    "}",
])

heading2("9.5  Cloud Deployment — Production")
data_table(
    ["Layer", "Platform", "Config file", "Key env var"],
    [
        ["Frontend",  "Vercel",      "frontend/vercel.json",     "VITE_API_URL — set to Render backend URL"],
        ["Backend",   "Render.com",  "backend/render.yaml",      "ALLOWED_ORIGINS — set to Vercel app URL"],
    ],
    col_widths=[1.0, 1.2, 2.2, 2.3]
)
body("vercel.json configures the Vite framework preset and SPA rewrites (all routes → /index.html).")
body("render.yaml configures: pip install -r requirements.txt; uvicorn main:app --host 0.0.0.0 --port $PORT.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  10. SECURITY DESIGN
# ═══════════════════════════════════════════════════════════════════════════════
heading1("10.  Security Design")

data_table(
    ["Control", "Implementation"],
    [
        ["CORS restriction",          "Backend restricts API access via ALLOWED_ORIGINS env var. In production set to exact Vercel origin — prevents cross-site abuse."],
        ["No persistent storage",     "No database, no server-side file storage. Uploaded files are read into memory, parsed, and discarded — nothing is written to disk."],
        ["File type validation",      "Only .xlsx and .xls files accepted. HTTP 400 returned for other extensions before any parsing occurs."],
        ["No credentials in frontend","There are no SAP credentials, API keys, or secrets in the React bundle. VITE_API_URL is the only env var passed to the frontend."],
        ["HTTPS enforced",            "Both Vercel and Render.com enforce HTTPS by default. No plain HTTP in production."],
        ["Error detail control",      "Traceback detail in HTTP 422 responses is useful for debugging but should be disabled or logged server-side in a production-hardened deployment."],
        ["localStorage scope",        "Upload history is stored in browser localStorage — scoped to the origin, not transmitted to the server. Contains full parse results, so advise against shared devices."],
    ],
    col_widths=[1.9, 4.8]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  11. CONFIGURATION & ENVIRONMENT VARIABLES
# ═══════════════════════════════════════════════════════════════════════════════
heading1("11.  Configuration & Environment Variables")

heading2("11.1  Backend (.env / environment)")
data_table(
    ["Variable", "Default", "Description"],
    [
        ["ALLOWED_ORIGINS",  "*",       "Comma-separated list of allowed CORS origins. Set to Vercel URL in production."],
        ["SAP_BASE_URL",     "(none)",  "Reserved — SAP Gateway base URL for future OData integration (sap_client.py)."],
        ["SAP_CLIENT",       "100",     "Reserved — SAP client number."],
        ["SAP_USER",         "(none)",  "Reserved — SAP service account username."],
        ["SAP_PASSWORD",     "(none)",  "Reserved — SAP service account password."],
        ["SAP_ODATA_SERVICE","(none)",  "Reserved — OData service path (e.g. /sap/opu/odata/sap/CNVLTVF3_SRV)."],
        ["USE_MOCK",         "true",    "Reserved — switches sap_client.py between mock and live SAP (not currently used by main.py)."],
    ],
    col_widths=[2.0, 1.2, 3.5]
)

heading2("11.2  Frontend (Vite env vars)")
data_table(
    ["Variable", "Set in", "Description"],
    [
        ["VITE_API_URL", "Vercel dashboard / .env.local", "Backend base URL. If unset, sapApi.ts uses empty string (relies on Vite proxy in dev or nginx proxy in Docker)."],
    ],
    col_widths=[1.6, 2.2, 2.9]
)

heading2("11.3  vite.config.ts — Dev Proxy")
body(
    "In development, all requests starting with /api are proxied to http://localhost:8000 "
    "by Vite's built-in proxy. This means the frontend can call /api/upload without CORS "
    "issues and without needing VITE_API_URL set."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  12. LOCAL DEVELOPMENT SETUP
# ═══════════════════════════════════════════════════════════════════════════════
heading1("12.  Local Development Setup")

heading2("12.1  Prerequisites")
for item in [
    "Docker Desktop (for Docker Compose workflow)",
    "Node.js 20+ and npm (for running frontend without Docker)",
    "Python 3.12+ and pip (for running backend without Docker)",
]:
    bullet(item)

heading2("12.2  Option A — Docker Compose (recommended)")
code_block([
    "cd C:\\Users\\I750407\\ltvf-dashboard",
    "",
    "# Copy env template (only needed once)",
    "copy backend\\.env.example backend\\.env",
    "",
    "# Build and start both services",
    "docker-compose up --build",
    "",
    "# Open in browser",
    "# http://localhost:3000",
])

heading2("12.3  Option B — Manual (no Docker)")
code_block([
    "# Terminal 1 — Backend",
    "cd backend",
    "pip install -r requirements.txt",
    "copy .env.example .env",
    "uvicorn main:app --reload --port 8000",
    "",
    "# Terminal 2 — Frontend",
    "cd frontend",
    "npm install",
    "npm run dev",
    "# Opens at http://localhost:3000",
])

heading2("12.4  Running Tests / Build Check")
code_block([
    "# Type-check the frontend",
    "cd frontend && npx tsc --noEmit",
    "",
    "# Production build (checks for build errors)",
    "cd frontend && npm run build",
    "",
    "# FastAPI docs (confirm backend is running)",
    "# http://localhost:8000/docs",
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  13. PRODUCTION DEPLOYMENT
# ═══════════════════════════════════════════════════════════════════════════════
heading1("13.  Production Deployment")

heading2("13.1  Deploy Backend to Render.com")
for step in [
    "Push the repository to GitHub.",
    "Create a new Render Web Service — connect to the repository, set root directory to backend/.",
    "Build command: pip install -r requirements.txt",
    "Start command:  uvicorn main:app --host 0.0.0.0 --port $PORT",
    "Add environment variable: ALLOWED_ORIGINS = https://<your-vercel-app>.vercel.app",
    "Note the Render service URL (e.g. https://ltvf-backend.onrender.com).",
]:
    bullet(step)

heading2("13.2  Deploy Frontend to Vercel")
for step in [
    "Import the repository into Vercel — set root directory to frontend/.",
    "Framework preset: Vite (auto-detected from vercel.json).",
    "Add environment variable: VITE_API_URL = https://ltvf-backend.onrender.com",
    "Deploy — Vercel builds with npm run build and serves the dist/ folder.",
    "SPA rewrites in vercel.json route all paths to /index.html.",
]:
    bullet(step)

info_box(
    "Production Checklist",
    [
        "  [ ]  ALLOWED_ORIGINS on Render set to exact Vercel URL (no wildcard in production)",
        "  [ ]  VITE_API_URL on Vercel set to exact Render URL",
        "  [ ]  Both URLs use HTTPS",
        "  [ ]  Upload a real CNVLTVF3 .xlsx export and verify parsing",
        "  [ ]  Check /api/health returns 200 OK",
    ],
    color_hex="E6F4EA",
    border_hex="16A34A"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  14. GLOSSARY
# ═══════════════════════════════════════════════════════════════════════════════
heading1("14.  Glossary")

data_table(
    ["Term", "Definition"],
    [
        ["LTVF",           "Long-Term Value Forecast — SAP test framework for comparing source and target data during system conversions."],
        ["CNVLTVF3",       "SAP transaction code for the LTVF test framework. Produces the .xlsx export parsed by this dashboard."],
        ["Match rate",     "Percentage of records that match exactly between source and target (0–100%). Core metric in CNVLTVF3."],
        ["Pass",           "Test case with match rate >= 95% (configurable). Displayed in green."],
        ["Warn",           "Test case with 80% <= match rate < 95% (configurable). Displayed in amber."],
        ["Fail",           "Test case with match rate < 80% (configurable). Displayed in red."],
        ["Group row",      "A hierarchy header row in the Excel file — identified by a '>>' prefix in the test_name column."],
        ["Leaf row",       "An actual test case row (no '>>' prefix). Only leaf rows contribute to pass/warn/fail counts."],
        ["SPA",            "Single-Page Application — loads once, updates dynamically. All navigation is client-side."],
        ["FastAPI",        "Python web framework used for the backend REST API. Provides automatic OpenAPI documentation."],
        ["Pydantic",       "Python data validation library used to define and enforce request/response schemas."],
        ["Vite",           "Frontend build tool and dev server. Handles TypeScript compilation, Tailwind, and hot module reload."],
        ["AG Grid",        "High-performance browser data grid used for the Detail Table tab."],
        ["Recharts",       "React charting library used for all charts: PieChart, BarChart, Treemap."],
        ["Tailwind CSS",   "Utility-first CSS framework. darkMode: 'class' strategy enables theme switching."],
        ["Render.com",     "Cloud platform hosting the FastAPI backend. Free tier available; render.yaml drives deployment."],
        ["Vercel",         "Cloud platform hosting the React frontend. vercel.json drives deployment and SPA rewrites."],
        ["localStorage",   "Browser-side key-value store used to persist theme preference and upload history (max 5 entries)."],
        ["OData",          "SAP standard REST protocol for exposing SAP data as JSON/XML. Reserved for future sap_client.py integration."],
        ["CORS",           "Cross-Origin Resource Sharing — browser security policy. ALLOWED_ORIGINS restricts which origins can call the backend API."],
    ],
    col_widths=[1.5, 5.2]
)

# ── Footer ────────────────────────────────────────────────────────────────────
doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer_p.add_run(
    f"LTVF Cloud Dashboard — Technical Reference Document  |  "
    f"Version 2.0  |  {datetime.date.today().strftime('%d %B %Y')}  |  Business Confidential"
)
r.font.size = Pt(8)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
r.font.italic = True

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = "C:/Users/I750407/ltvf-dashboard/LTVF_Technical_Reference.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
