from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Margins ───────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(3.0)

# ── Colours ───────────────────────────────────────────────────────────────────
SAP_BLUE  = RGBColor(0x00, 0x33, 0x66)
MID_BLUE  = RGBColor(0x00, 0x5A, 0x9C)
DARK      = RGBColor(0x22, 0x22, 0x22)
GREY      = RGBColor(0x55, 0x55, 0x55)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
AMBER     = RGBColor(0xB4, 0x5A, 0x00)
GREEN     = RGBColor(0x15, 0x6B, 0x37)
RED       = RGBColor(0x99, 0x1B, 0x1B)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = SAP_BLUE
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "003366")
    pBdr.append(bot)
    pPr.append(pBdr)


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = MID_BLUE


def body(text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = color or DARK


def bullet(text, level=0, color=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.7 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = color or DARK


def note_box(label, text, bg="FFF8E1", border="B45A00", label_color=None):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBdr = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    "6")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), border)
        tcBdr.append(b)
    tcPr.append(tcBdr)
    cell.paragraphs[0].clear()
    lr = cell.paragraphs[0].add_run(label + "  ")
    lr.font.bold = True
    lr.font.size = Pt(10.5)
    lr.font.color.rgb = label_color or AMBER
    tr = cell.paragraphs[0].add_run(text)
    tr.font.size = Pt(10.5)
    tr.font.color.rgb = DARK
    doc.add_paragraph()


def simple_table(headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hrow = tbl.rows[0]
    for i, h in enumerate(headers):
        c = hrow.cells[i]
        set_cell_bg(c, "003366")
        c.paragraphs[0].clear()
        r = c.paragraphs[0].add_run(h)
        r.font.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        tr = tbl.add_row()
        bg = "F5F5F5" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            c = tr.cells[ci]
            set_cell_bg(c, bg)
            tc = c._tc
            tcPr = tc.get_or_add_tcPr()
            tcBdr = OxmlElement("w:tcBorders")
            for side in ("top","left","bottom","right"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"),   "single")
                b.set(qn("w:sz"),    "4")
                b.set(qn("w:space"), "0")
                b.set(qn("w:color"), "CCCCCC")
                tcBdr.append(b)
            tcPr.append(tcBdr)
            c.paragraphs[0].clear()
            r = c.paragraphs[0].add_run(val)
            r.font.size = Pt(10)
            r.font.color.rgb = DARK
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
# COVER
# ═══════════════════════════════════════════════════════════════════════════════
banner = doc.add_table(rows=1, cols=1)
banner.style = "Table Grid"
bc = banner.cell(0, 0)
set_cell_bg(bc, "003366")
bc.paragraphs[0].clear()
bc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = bc.paragraphs[0].add_run("\n\nLTVF Cloud Dashboard\n")
r1.font.size = Pt(26)
r1.font.bold = True
r1.font.color.rgb = WHITE
p2 = bc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Ideation & Architecture Discussion\n\n")
r2.font.size = Pt(13)
r2.font.color.rgb = RGBColor(0xB3, 0xD1, 0xFF)
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in [
    f"Date: {datetime.date.today().strftime('%d %B %Y')}",
    "Status: DRAFT — For Discussion Only",
    "Audience: Architecture Review / Business Stakeholders",
]:
    r = meta.add_run(line + "\n")
    r.font.size = Pt(10.5)
    r.font.color.rgb = GREY
    r.font.italic = True

note_box(
    "Note:",
    "This is a rough ideation document intended to spark architecture discussion and gather "
    "initial feedback. It is deliberately high-level. Details will be refined based on input "
    "from this review.",
    bg="FFF3CD", border="B45A00"
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 1. THE IDEA — WHAT ARE WE TRYING TO DO?
# ═══════════════════════════════════════════════════════════════════════════════
h1("1.  The Idea — What Are We Trying to Do?")

body(
    "Right now, the LTVF report lives entirely inside SAP GUI. To view it, a user needs "
    "SAP GUI installed, a valid SAP login, and usually VPN access. "
    "That creates friction — especially for finance managers, business analysts, or external "
    "stakeholders who only need to read the report, not operate SAP."
)

body(
    "The idea is simple: take the same LTVF data and surface it through a modern web "
    "browser — no SAP GUI, no VPN dependency, no special software. The data stays in SAP. "
    "We just put a better window in front of it."
)

note_box(
    "In one sentence:",
    "A cloud-hosted dashboard that reads LTVF data from SAP and displays it in a browser — "
    "accessible to anyone with the URL and credentials.",
    bg="E8F0FE", border="003366", label_color=SAP_BLUE
)

h2("What the current SAP screen looks like vs what we're building")

simple_table(
    ["", "SAP GUI (Today)", "Cloud Dashboard (Proposed)"],
    [
        ["Access",       "SAP Logon + VPN",              "Any browser, any device"],
        ["Who can see it", "SAP-licensed users only",     "Any authorised business user"],
        ["Layout",       "Fixed ALV grid",               "Interactive table + KPI cards"],
        ["Export",       "Manual SAP list export",       "One-click download"],
        ["Sharing",      "Screenshot or export file",    "Share a URL"],
    ],
    col_widths=[1.3, 2.5, 2.5]
)


# ═══════════════════════════════════════════════════════════════════════════════
# 2. HIGH-LEVEL ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════════
h1("2.  High-Level Architecture")

body(
    "Three layers — each doing one job:"
)

p = doc.add_paragraph()
r = p.add_run(
    "  [ Browser ]  ←→  [ Cloud API ]  ←→  [ SAP S/4HANA ]\n"
    "   React UI        FastAPI (Python)      OData / Gateway"
)
r.font.name = "Courier New"
r.font.size = Pt(10)
r.font.bold = True
r.font.color.rgb = SAP_BLUE
doc.add_paragraph()

bullet("Browser (React):  renders the report table, KPI numbers, and charts. Talks only to our API — never directly to SAP.")
bullet("Cloud API (FastAPI):  the middleman. Fetches data from SAP, formats it, hands it to the browser. Holds SAP credentials securely.")
bullet("SAP:  the system of record. Nothing changes here. We read data through the existing SAP Gateway OData interface.")

note_box(
    "Key point:",
    "SAP is untouched. We are building a read-only view on top of existing data. "
    "There is no risk of data modification and no SAP development required (assuming an OData service exists).",
    bg="E6F4EA", border="156B37", label_color=GREEN
)


# ═══════════════════════════════════════════════════════════════════════════════
# 3. FRONTEND — REACT
# ═══════════════════════════════════════════════════════════════════════════════
h1("3.  Frontend — Why React?")

h2("What the frontend does")
bullet("Displays the LTVF hierarchy as an interactive table — matching the SAP layout but better.")
bullet("Shows KPI summary cards at the top: Net MRP, Global Amount, Delta.")
bullet("Colour-codes positive/negative deltas (green / red) — same as SAP traffic lights.")
bullet("Refresh button to pull latest data from SAP on demand.")
doc.add_paragraph()

h2("Why React — and not something else?")

body(
    "This came up early. Below is honest reasoning, not marketing:"
)

simple_table(
    ["Option", "Why we looked at it", "Why we moved on"],
    [
        ["React",           "Largest ecosystem, best AG Grid support, huge talent pool",            "— Selected —"],
        ["Next.js",         "React-based, great for SEO and server-side rendering",                 "SSR not needed for an internal dashboard. Adds complexity we don't need yet."],
        ["Angular",         "Enterprise-grade, Google-backed, strong TypeScript",                   "Heavy framework for a read-only report. Steeper learning curve, larger bundle."],
        ["Vue 3",           "Lightweight, gentle learning curve, good for small apps",              "AG Grid enterprise integration is less mature than React. Smaller hiring pool."],
        ["Svelte / SvelteKit", "Very fast, minimal boilerplate, growing community",                 "Immature ecosystem for data-heavy grid components. Higher risk for enterprise."],
        ["Plain HTML + JS", "No framework overhead, fastest to ship a prototype",                   "Becomes unmaintainable quickly as features grow. No component reuse."],
    ],
    col_widths=[1.4, 2.5, 2.5]
)

note_box(
    "The real reason:",
    "The LTVF report is a deeply hierarchical data grid. AG Grid — which has the best React "
    "integration — handles tree data, large row counts, sorting, and filtering out of the box. "
    "That alone makes React the practical choice for this specific use case.",
    bg="E8F0FE", border="003366", label_color=SAP_BLUE
)

h2("What the screen looks like (rough layout)")

p = doc.add_paragraph()
r = p.add_run(
    "  ┌──────────────────────────────────────────────────────────┐\n"
    "  │  LTVF Results: BOV-ENC_SDT_TC1_FI    System: QSL  [↻]  │  ← Header\n"
    "  ├────────────┬─────────────┬────────────┬──────────────────┤\n"
    "  │ Net MRP    │ Global Amt  │   Delta    │   Line Items     │  ← KPI cards\n"
    "  │ 1,129,229  │ 1,622,064   │  -492,835  │      14          │\n"
    "  ├────────────┴─────────────┴────────────┴──────────────────┤\n"
    "  │ Description          │ Total │ Net MRP │ Delta  │ Δ %    │  ← AG Grid\n"
    "  │ ▶ Tool Costs         │ 2,576 │ 1,129K  │ -492K  │ -43.7% │    table\n"
    "  │   ▶ MAESTR DATA 1.1  │   319 │   108K  │      0 │   0.0% │\n"
    "  │     ROBOT AGREEMENT  │   319 │   108K  │      0 │   0.0% │\n"
    "  │ ▶ Business Partners  │ 4,716 │   984K  │ -423K  │ -43.0% │\n"
    "  └──────────────────────────────────────────────────────────┘\n"
)
r.font.name = "Courier New"
r.font.size = Pt(8.5)
r.font.color.rgb = DARK
doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
# 4. BACKEND — FASTAPI
# ═══════════════════════════════════════════════════════════════════════════════
h1("4.  Backend — FastAPI (Python)")

h2("What the backend does")
bullet("Acts as the secure bridge between the browser and SAP.")
bullet("Holds SAP credentials — the browser never sees them.")
bullet("Calls the SAP OData service, normalises the JSON response, returns it to the frontend.")
bullet("Provides a mock mode — returns sample data when SAP is not yet connected.")
doc.add_paragraph()

h2("Why FastAPI over other backend options?")

simple_table(
    ["Option", "Considered?", "Verdict"],
    [
        ["FastAPI (Python)",   "Yes", "Selected. Fast to build, auto-generates API docs, clean SAP OData integration via requests library."],
        ["Node.js / Express",  "Yes", "Viable. But Python has more mature SAP integration libraries and the team is more comfortable with it."],
        ["Java Spring Boot",   "Yes", "Technically strong but heavy. Overkill for a read-only reporting API. Slower iteration."],
        ["SAP BTP / CAP",      "Yes", "Native SAP cloud platform. Keeps everything in SAP ecosystem but locks us in and adds cost."],
        ["Serverless (Lambda/Functions)", "Briefly", "Could work for low-traffic. Cold start latency is a concern for a dashboard that needs to feel responsive."],
    ],
    col_widths=[1.8, 1.0, 4.0]
)

h2("SAP Connection approach")
bullet("Protocol: SAP OData via HTTPS — standard, supported, no custom SAP development needed.")
bullet("Auth: HTTP Basic Auth with a dedicated read-only service account (not a personal login).")
bullet("Mock toggle: USE_MOCK=true/false — one environment variable switches between live SAP and sample data.")

note_box(
    "Still to confirm with BASIS team:",
    "We need the OData service URL, HTTPS port, and a service account before we can go live. "
    "The architecture is ready — it's a config change, not a code change.",
    bg="FFF8E1", border="B45A00"
)


# ═══════════════════════════════════════════════════════════════════════════════
# 5. CHALLENGES — CUSTOMER PERSPECTIVE
# ═══════════════════════════════════════════════════════════════════════════════
h1("5.  Challenges — What Are We Worried About?")

body(
    "Being honest here. These are the real things that could cause friction "
    "and they are worth discussing before we go further."
)

h2("5.1  SAP Connectivity")
bullet("Risk: SAP Gateway OData may not be enabled, or the LTVF program data may not be exposed as an OData service.")
bullet("Impact: If no OData service exists, BASIS would need to create one — adding time and SAP transport effort.")
bullet("Mitigation: We are already running in mock mode. The frontend is fully built. The SAP connection is the last mile.")
doc.add_paragraph()

h2("5.2  Data Accuracy & Trust")
bullet("Risk: Business users may not trust a web dashboard to show the same numbers as SAP GUI.")
bullet("Impact: Adoption could be slow if users keep double-checking against SAP.")
bullet("Mitigation: Add a 'Last refreshed' timestamp and a direct link back to the SAP report for validation during rollout.")
doc.add_paragraph()

h2("5.3  User Adoption")
bullet("Risk: Users are accustomed to the SAP ALV layout. A web UI — even a better one — requires a change in habit.")
bullet("Impact: Training overhead; resistance from power users who know SAP shortcuts.")
bullet("Mitigation: Keep the column names, row structure, and colour coding as close to SAP as possible. Run a parallel period.")
doc.add_paragraph()

h2("5.4  Authentication & Access Control")
bullet("Risk: The current build has no user login on the web app itself.")
bullet("Impact: Anyone with the URL can see the financial data — not acceptable for production.")
bullet("Mitigation: Must add authentication before go-live. Options: Azure AD SSO, Google OAuth, or a simple API key for internal use. This is a Phase 4 item.")
doc.add_paragraph()

h2("5.5  SAP Performance")
bullet("Risk: The OData call to SAP could be slow if the LTVF dataset is large or the Gateway is under load.")
bullet("Impact: The dashboard may feel slow to load — poor first impression.")
bullet("Mitigation: Add a loading spinner (already in place). Consider caching the response for 5–10 minutes if data does not need to be real-time.")
doc.add_paragraph()

h2("5.6  Network / Firewall")
bullet("Risk: The cloud-hosted backend may not be able to reach the SAP system on-premise due to firewall rules.")
bullet("Impact: Entire integration blocked until network path is opened.")
bullet("Mitigation: Raise with IT/BASIS early. If cloud-to-SAP is blocked, the backend can be deployed inside the corporate network instead.")
doc.add_paragraph()

h2("5.7  Keeping Up With SAP Changes")
bullet("Risk: If the LTVF program or OData service changes (e.g. due to an SAP upgrade or a CNV re-configuration), the dashboard could break.")
bullet("Impact: Silent failures — wrong or missing data.")
bullet("Mitigation: Add a health check endpoint and monitoring alert. Any OData schema change should trigger a dashboard review.")
doc.add_paragraph()

note_box(
    "Discussion question for this review:",
    "Which of these challenges is the biggest blocker from the business / customer side? "
    "Prioritising the mitigation order will shape the next phase of work.",
    bg="F3E8FF", border="6B21A8", label_color=RGBColor(0x6B, 0x21, 0xA8)
)


# ═══════════════════════════════════════════════════════════════════════════════
# 6. OPEN QUESTIONS
# ═══════════════════════════════════════════════════════════════════════════════
h1("6.  Open Questions for This Discussion")

for q in [
    "1.  Should the dashboard be accessible externally (internet) or internal network only?",
    "2.  Who are the primary users — finance team, management, external auditors?",
    "3.  Do we need user-level access control (some users see some cost centres, others see all)?",
    "4.  Is real-time data a hard requirement, or is a 10-minute cache acceptable?",
    "5.  Will the report need to support multiple SAP clients / systems (QSL and PRD)?",
    "6.  Is there a preference for cloud provider — Azure, AWS, GCP, or on-premise?",
    "7.  What does success look like for Phase 1 go-live — number of users, data scope?",
]:
    bullet(q)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# 7. WHAT IS BUILT TODAY
# ═══════════════════════════════════════════════════════════════════════════════
h1("7.  What Is Already Built")

body("To give this discussion something concrete — here is where we are today:")

simple_table(
    ["Component", "Status", "Detail"],
    [
        ["React frontend",     "Done",        "Running at localhost:3000. Full LTVF table, KPI cards, header."],
        ["FastAPI backend",    "Done",        "Running at localhost:8000. Mock LTVF data. /api/ltvf endpoint live."],
        ["SAP OData client",   "Ready",       "Code written. Waiting for SAP Gateway URL and service account from BASIS."],
        ["Docker containers",  "Done",        "docker-compose.yml ready. One command to run both services."],
        ["Authentication",     "Not started", "Needed before production. Azure AD SSO recommended."],
        ["Cloud deployment",   "Not started", "Pending cloud platform decision."],
    ],
    col_widths=[1.8, 1.1, 3.9]
)

note_box(
    "Demo available:",
    "The mock-mode dashboard is running and can be demonstrated right now. "
    "It shows the LTVF report structure with realistic data — no SAP connection needed for the demo.",
    bg="E6F4EA", border="156B37", label_color=GREEN
)

# ── Footer ────────────────────────────────────────────────────────────────────
doc.add_paragraph()
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run(
    f"LTVF Cloud Dashboard — Ideation Draft  |  "
    f"{datetime.date.today().strftime('%d %B %Y')}  |  For Discussion Only"
)
r.font.size = Pt(8)
r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
r.font.italic = True

# ── Save ──────────────────────────────────────────────────────────────────────
out = "C:/Users/I750407/ltvf-dashboard/LTVF_Ideation_Draft.docx"
doc.save(out)
print(f"Saved: {out}")
